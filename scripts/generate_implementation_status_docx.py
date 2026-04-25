from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = REPO_ROOT / "docs"


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def build_document() -> Document:
    doc = Document()
    doc.add_heading("FinGuard Implementation Status", level=0)
    doc.add_paragraph(
        "This report summarizes the current implementation against the AAS practice-module expectations after the wrap-up work."
    )
    doc.add_paragraph(
        f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )

    doc.add_heading("1. Runtime completeness", level=1)
    _table(
        doc,
        ["Area", "Status", "Notes"],
        [
            ["Frontend demo shell", "Implemented", "Backed by current FastAPI APIs"],
            ["Portfolio analysis", "Implemented", "Returns crew_output plus analysis_trace"],
            ["Transaction -> case workflow", "Implemented", "High-risk transactions auto-open cases"],
            ["SAR export", "Implemented", "JSON and PDF endpoints available"],
            ["Deterministic demo seed", "Implemented", "scripts/seed_demo_data.py"],
            ["Deterministic AI mock mode", "Implemented", "AI_RESPONSE_MODE=mock"],
        ],
    )

    doc.add_heading("2. Assessment artifacts", level=1)
    _table(
        doc,
        ["Artifact", "Status", "Source"],
        [
            ["System architecture", "Implemented", "docs/SYSTEM_ARCHITECTURE.md"],
            ["Agent design", "Implemented", "docs/AGENT_DESIGN.md"],
            ["Responsible AI report", "Implemented", "docs/RESPONSIBLE_AI_REPORT.md"],
            ["AI security risk register", "Implemented", "docs/AI_SECURITY_RISK_REGISTER.md"],
            ["MLSecOps / LLMSecOps design", "Implemented", "docs/MLSECOPS_LLMSecOps_PIPELINE.md"],
            ["Testing summary / demo runbook", "Implemented", "docs/TESTING_AND_DEMO_RUNBOOK.md"],
        ],
    )

    doc.add_heading("3. Verification evidence", level=1)
    for line in (
        "Automated tests: pytest -q",
        "CI workflow: .github/workflows/ci.yml",
        "Docker health checks: docker-compose.yml",
        "Report source pack: docs/REPORT_SOURCE_PACK.md",
    ):
        doc.add_paragraph(line, style="List Bullet")

    return doc


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    output = DOCS_DIR / "FinGuard_Implementation_Status_Report.docx"
    build_document().save(output)
    print(f"Wrote {output}")


if __name__ == "__main__":
    main()

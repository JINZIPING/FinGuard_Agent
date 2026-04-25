from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = REPO_ROOT / "docs"


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value


def build_document() -> Document:
    doc = Document()
    doc.add_heading("FinGuard Assessment Summary", level=0)
    doc.add_paragraph(
        "This generated summary reflects the current FastAPI + LangGraph implementation and the assessment-pack artifacts stored in the repository."
    )

    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    _kv_table(
        doc,
        [
            ("Generated", generated_at),
            ("Frontend", "React shell frontend served by Nginx"),
            ("Backend", "FastAPI + SQLite + audit / case / SAR services"),
            ("AI system", "FastAPI + LangGraph + internal agents"),
            ("LLM mode", "OpenAI live mode or AI_RESPONSE_MODE=mock"),
        ],
    )

    doc.add_heading("1. Core capabilities", level=1)
    for line in (
        "Portfolio analysis with analysis_trace",
        "Transaction risk scoring with alert and auto-case generation",
        "Case review, customer 360, audit verification, and SAR export",
        "Deterministic demo seed workflow",
        "Automated tests and Docker smoke checks",
    ):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("2. Assessment artifacts", level=1)
    for line in (
        "docs/SYSTEM_ARCHITECTURE.md",
        "docs/AGENT_DESIGN.md",
        "docs/RESPONSIBLE_AI_REPORT.md",
        "docs/AI_SECURITY_RISK_REGISTER.md",
        "docs/MLSECOPS_LLMSecOps_PIPELINE.md",
        "docs/TESTING_AND_DEMO_RUNBOOK.md",
        "docs/REPORT_SOURCE_PACK.md",
    ):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("3. Demo controls", level=1)
    _kv_table(
        doc,
        [
            ("Seed data", "python scripts/seed_demo_data.py --reset"),
            ("Mock mode", "AI_RESPONSE_MODE=mock"),
            ("Automated tests", "pytest -q"),
            ("Local stack", "docker compose up --build"),
        ],
    )

    return doc


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    output = DOCS_DIR / "FinGuard_Assessment_Summary.docx"
    build_document().save(output)
    print(f"Wrote {output}")


if __name__ == "__main__":
    main()
